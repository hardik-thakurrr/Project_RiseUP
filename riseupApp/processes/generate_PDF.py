

import pandas as pd
from docx import Document
from datetime import datetime
from docx import Document
from docx.shared import Inches,Cm
from datetime import datetime
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt
from docx.shared import RGBColor
from .responseValidator import *

import pythoncom
from docx2pdf import convert
import os

def add_formatted_section(doc,title, content):
    if title:
        heading = doc.add_paragraph()
        run = heading.add_run(title)
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(0, 0, 0)
        run.font.bold = True
        heading.paragraph_format.space_before = Pt(0)
        heading.paragraph_format.space_after = Pt(6)
   
    for line in content.split('\n'):
        stripped_line = line.strip()
        if stripped_line:
            if stripped_line.endswith(':'):
                doc.add_paragraph(stripped_line)
            else:
                doc.add_paragraph(stripped_line, style="List Bullet")

# Function to generate the PDF
def generate_PDF(json_data, reportDirPath, userName):

    pythoncom.CoInitialize()

    # Creating a DataFrame from the JSON data
    df=pd.DataFrame(json_data).T

    total_points = df["Score"].str.extract(r'(\d)').astype(int).sum().item()
    print(total_points)

    
    overall_summary = CandidateSummary(df)
    
    # Create a document
    doc = Document()

    # Access the section and its header
    section = doc.sections[0]
    header = section.header

    # Clear existing header content (optional)
    if header.paragraphs:
        for para in header.paragraphs:
            para.clear()

    # Create a table in the header (1 row, 2 columns)
    table = header.add_table(rows=1, cols=2,width=Cm(16.5))  # 16.5 cm = 6.5 inches
    table.autofit = False  # Disable autofit so we can set custom widths

    # Set column widths
    table.columns[0].width = Inches(0.1)  # Space for the image
    table.columns[1].width = Inches(1)  # Space for the date

    # Add the image to the left cell
    left_cell = table.cell(0, 0)
    left_para = left_cell.paragraphs[0] if left_cell.paragraphs else left_cell.add_paragraph()
    run = left_para.add_run()

    script_dir = os.path.dirname(os.path.abspath(__file__))  # Directory of the current script
    image_path = os.path.join(script_dir, 'RiseUP.png')
    run.add_picture(image_path, height=Cm(2))

    # Add the date to the right cell
    right_cell = table.cell(0, 1)
    right_para = right_cell.paragraphs[0] if right_cell.paragraphs else right_cell.add_paragraph()
    right_para.alignment = 2  # Right-align text
    right_para.add_run(datetime.now().strftime("%d-%m-%Y"))  # Current date



    # Add a Title to the Document
    heading = doc.add_heading('Interview Evaluation', level=1)

    # Style the heading
    run = heading.runs[0]
    run.font.size = Pt(24)
    run.font.bold = True
    run.font.underline = True
    run.font.color.rgb = RGBColor(0, 0, 0)

    # Center align the heading
    heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    # Adjust spacing before and after
    heading.paragraph_format.space_before = Pt(0)
    heading.paragraph_format.space_after = Pt(6)

    # Add  Name to the Document
    Name = doc.add_heading(userName, level=1)

    # Style the heading
    run = Name.runs[0]
    run.font.size = Pt(22)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 0, 255)

    # Center align the heading
    Name.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    # Adjust spacing before and after
    Name.paragraph_format.space_before = Pt(0)
    Name.paragraph_format.space_after = Pt(6)

    # Add Total Points as a regular paragraph
    point = doc.add_paragraph()
    run = point.add_run('Total Points: ')
    # Style the text
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0, 0, 0)

    # Determine the score color based on variable_marks
    if total_points < 7:
        score_color = RGBColor(255, 0, 0)  # Red
    elif total_points < 14:
        score_color = RGBColor(255, 165, 0)  # Orange
    else:
        score_color = RGBColor(0, 128, 0)  # Green

    # Add the score in the determined color
    score_run = point.add_run(f'{total_points}/20')
    score_run.font.size = Pt(16)
    score_run.font.color.rgb = score_color

    # Center align the paragraph
    point.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Adjust spacing before and after
    point.paragraph_format.space_before = Pt(0)
    point.paragraph_format.space_after = Pt(6)

    # Add Summary  as a regular paragraph
    summary = doc.add_paragraph()
    run = summary.add_run('Overall Summary')

    # Style the text
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0, 0, 0)

    # Center align the paragraph
    summary.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Adjust spacing before and after
    summary.paragraph_format.space_before = Pt(0)
    summary.paragraph_format.space_after = Pt(6)

    # Add Summary_Content  as a regular paragraph
    summary_c = doc.add_paragraph()
    run = summary_c.add_run(overall_summary)

    # Style the text
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0, 0, 0)

    # Center align the paragraph
    summary.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Adjust spacing before and after
    summary.paragraph_format.space_before = Pt(0)
    summary.paragraph_format.space_after = Pt(6)


    # Add Question Wise Evaluation  as a regular paragraph
    summary = doc.add_paragraph()
    run = summary.add_run('Question Wise Evaluation')

    # Style the text
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0, 0, 0)

    # Center align the paragraph
    summary.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Adjust spacing before and after
    summary.paragraph_format.space_before = Pt(0)
    summary.paragraph_format.space_after = Pt(6)

    # # Iterate over DataFrame and add formatted content
    for index, row in df.iterrows():
        doc.add_heading(f"{index}:", level=1)
        question =doc.add_paragraph()
        run = question.add_run(row['question'] if pd.notna(row['question']) else "")
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(0, 0, 0)
        run.font.italic = True
        

        score = doc.add_paragraph()
        run_black = score.add_run('Score: ')
        run_black.font.size = Pt(14)
        run_black.font.color.rgb = RGBColor(0, 0, 0)
        run_black.font.bold = True

    
        score_value = row['Score'] if pd.notna(row['Score']) else ""  # Handle NaN values
        if score_value:  # Check if the value is not empty
            score_value = score_value.strip()  # Remove leading/trailing spaces
            if "points" not in score_value:  # Append "point" if not already present
                score_value += " points"

        run_blue = score.add_run(score_value)
        run_blue.font.size = Pt(14)
        run_blue.font.color.rgb = RGBColor(0, 0, 255)
        run_blue.font.bold = True

        

        add_formatted_section(doc,"Candidate Response:",row['candidateResponse'] if pd.notna(row['candidateResponse']) else "")
        add_formatted_section(doc,"Strength:",row['Strengths'] if pd.notna(row['Strengths']) else "")
        add_formatted_section(doc,"Areas for Improvement:",row['AreasForImprovement'] if pd.notna(row['AreasForImprovement']) else "")
        add_formatted_section(doc,"Scoring Rubrics:",row['Rubric'] if pd.notna(row['Rubric']) else "")
    
    os.makedirs(reportDirPath, exist_ok=True)  # Create directory if it doesn't exist
    reportPath = os.path.join(reportDirPath, "Interview_Assessment.docx")

    doc.save(reportPath)
    print("Doc successfully created!")
    # Converting docx present in the same folder
    # as the python file
    convert(reportPath)

    return total_points, overall_summary




